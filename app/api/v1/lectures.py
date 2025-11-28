"""API эндпоинты для работы с лекциями"""
import os
import shutil
import logging
import subprocess
from datetime import datetime
from pathlib import Path
from typing import List, Optional
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, BackgroundTasks, Request
from fastapi.responses import JSONResponse, StreamingResponse
import json
import threading
from sqlalchemy.orm import Session
from sqlalchemy import create_engine
from app.core.database import engine

logger = logging.getLogger(__name__)

# Импорт модуля транскрибации
from app.utils.transcription import (
    transcribe_file, 
    find_ffmpeg, 
    ensure_ffmpeg_in_path,
    WHISPER_AVAILABLE
)

from app.core.database import get_db
from app.core.security import get_current_user
from app.core.limiter import limiter
from app.api.v1.dependencies import (
    require_course_access,
    require_lecture_access,
    require_lecture_teacher_access,
    require_material_access
)
from app.models import Course, Lecture, LectureMaterial, ProcessedMaterial, User, Test, Question
from app.schemas import CreateLectureRequest, UpdateLectureRequest, LectureMaterialResponse, LectureResponse, TestResponse, QuestionResponse

router = APIRouter()

# Директория для хранения файлов лекций
UPLOAD_DIR = Path("uploads/lectures")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

# Импорты для парсинга файлов
try:
    import PyPDF2
    PDF_PARSER_AVAILABLE = True
except ImportError:
    PDF_PARSER_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document
    DOCX_PARSER_AVAILABLE = True
except ImportError:
    DOCX_PARSER_AVAILABLE = False


@router.get("/courses/{course_id}/lectures", response_model=List[LectureResponse])
def get_course_lectures(
    course_id: int,
    course: Course = Depends(require_course_access),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Получение всех лекций курса"""
    # Проверка доступа выполнена через зависимость require_course_access
    
    # Получаем лекции в зависимости от роли с предзагрузкой материалов (избегаем N+1)
    from sqlalchemy.orm import selectinload
    if current_user.role == "admin" or current_user.role == "teacher":
        # Админ и преподаватель видят все лекции курса
        lectures = db.query(Lecture).options(
            selectinload(Lecture.materials)
        ).filter(Lecture.course_id == course_id).all()
    elif current_user.role == "student":
        # Студент видит только опубликованные лекции
        lectures = db.query(Lecture).options(
            selectinload(Lecture.materials)
        ).filter(
            Lecture.course_id == course_id,
            Lecture.published == True
        ).all()
    else:
        raise HTTPException(status_code=403, detail="Доступ запрещен")
    
    result = []
    for lecture in lectures:
        # Материалы уже загружены через selectinload, сортируем в памяти
        try:
            materials = sorted(lecture.materials, key=lambda m: getattr(m, 'order_index', 0))
        except Exception:
            # Если сортировка не удалась, используем исходный порядок
            materials = lecture.materials
        result.append(LectureResponse(
            id=lecture.id,
            course_id=lecture.course_id,
            name=lecture.name,
            description=lecture.description,
            created_at=lecture.created_at,
            published=lecture.published or False,
            generate_test=lecture.generate_test or False,
            test_generation_mode=lecture.test_generation_mode or "once",
            test_max_attempts=lecture.test_max_attempts or 1,
            test_show_answers=lecture.test_show_answers or False,
            test_deadline=lecture.test_deadline,
            materials=[LectureMaterialResponse(
                id=m.id,
                file_path=m.file_path,
                file_type=m.file_type,
                file_name=m.file_name,
                file_size=m.file_size,
                order_index=m.order_index
            ) for m in materials]
        ))
    
    return result


@router.post("/lectures", response_model=LectureResponse)
def create_lecture(
    payload: CreateLectureRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Создание новой лекции"""
    if current_user.role != "teacher":
        raise HTTPException(status_code=403, detail="Доступ разрешен только преподавателям")
    
    # Проверяем доступ к курсу через зависимость
    course = require_course_access(payload.course_id, db, current_user)
    
    # Создаем лекцию
    lecture = Lecture(
        course_id=payload.course_id,
        name=payload.name,
        description=payload.description,
        created_at=datetime.now().isoformat(),
        generate_test=payload.generate_test or False,
        test_generation_mode=payload.test_generation_mode or "once",
        test_max_attempts=payload.test_max_attempts or 1,
        test_show_answers=payload.test_show_answers or False,
        test_deadline=payload.test_deadline
    )
    
    db.add(lecture)
    db.commit()
    db.refresh(lecture)
    
    # При создании лекции материалов еще нет, возвращаем пустой список
    return LectureResponse(
        id=lecture.id,
        course_id=lecture.course_id,
        name=lecture.name,
        description=lecture.description,
        created_at=lecture.created_at,
        published=lecture.published or False,
        generate_test=lecture.generate_test or False,
        test_generation_mode=lecture.test_generation_mode or "once",
        test_max_attempts=lecture.test_max_attempts or 1,
        test_show_answers=lecture.test_show_answers or False,
        test_deadline=lecture.test_deadline,
        materials=[]
    )


@router.get("/lectures/{lecture_id}", response_model=LectureResponse)
def get_lecture(
    lecture_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Получение лекции по ID"""
    # Проверяем доступ через зависимость
    require_lecture_access(
        lecture_id, 
        db, 
        current_user, 
        require_published=(current_user.role == "student")
    )
    
    # Перезагружаем лекцию с предзагрузкой материалов (избегаем N+1)
    from sqlalchemy.orm import selectinload
    lecture = db.query(Lecture).options(
        selectinload(Lecture.materials)
    ).filter(Lecture.id == lecture_id).first()
    
    if not lecture:
        raise HTTPException(status_code=404, detail="Лекция не найдена")
    
    # Сортируем материалы в памяти
    try:
        materials = sorted(lecture.materials, key=lambda m: getattr(m, 'order_index', 0))
    except Exception:
        # Если сортировка не удалась, используем исходный порядок
        materials = lecture.materials
    
    return LectureResponse(
        id=lecture.id,
        course_id=lecture.course_id,
        name=lecture.name,
        description=lecture.description,
        created_at=lecture.created_at,
        published=lecture.published or False,
        generate_test=lecture.generate_test or False,
        test_generation_mode=lecture.test_generation_mode or "once",
        test_max_attempts=lecture.test_max_attempts or 1,
        test_show_answers=lecture.test_show_answers or False,
        test_deadline=lecture.test_deadline,
        materials=[{
            "id": m.id,
            "file_path": m.file_path,
            "file_type": m.file_type,
            "file_name": m.file_name,
            "file_size": m.file_size,
            "order_index": m.order_index
        } for m in materials]
    )


@router.put("/lectures/{lecture_id}")
def update_lecture(
    lecture_id: int,
    payload: UpdateLectureRequest,
    lecture: Lecture = Depends(require_lecture_teacher_access),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Обновление лекции"""
    # Проверка доступа выполнена через зависимость require_lecture_teacher_access
    
    # Обновляем поля
    if payload.name is not None:
        lecture.name = payload.name
    if payload.description is not None:
        lecture.description = payload.description
    if payload.generate_test is not None:
        lecture.generate_test = payload.generate_test
    if payload.test_generation_mode is not None:
        lecture.test_generation_mode = payload.test_generation_mode
    if payload.test_max_attempts is not None:
        lecture.test_max_attempts = payload.test_max_attempts
    if payload.test_show_answers is not None:
        lecture.test_show_answers = payload.test_show_answers
    if payload.test_deadline is not None:
        lecture.test_deadline = payload.test_deadline
    
    db.commit()
    
    # Перезагружаем лекцию с предзагрузкой материалов (избегаем N+1)
    from sqlalchemy.orm import selectinload
    lecture = db.query(Lecture).options(
        selectinload(Lecture.materials)
    ).filter(Lecture.id == lecture_id).first()
    
    # Сортируем материалы в памяти
    try:
        materials = sorted(lecture.materials, key=lambda m: getattr(m, 'order_index', 0))
    except Exception:
        # Если сортировка не удалась, используем исходный порядок
        materials = lecture.materials
    
    return LectureResponse(
        id=lecture.id,
        course_id=lecture.course_id,
        name=lecture.name,
        description=lecture.description,
        created_at=lecture.created_at,
        published=lecture.published or False,
        generate_test=lecture.generate_test or False,
        test_generation_mode=lecture.test_generation_mode or "once",
        test_max_attempts=lecture.test_max_attempts or 1,
        test_show_answers=lecture.test_show_answers or False,
        test_deadline=lecture.test_deadline,
        materials=[{
            "id": m.id,
            "file_path": m.file_path,
            "file_type": m.file_type,
            "file_name": m.file_name,
            "file_size": m.file_size,
            "order_index": m.order_index
        } for m in materials]
    )


@router.delete("/lectures/{lecture_id}")
def delete_lecture(
    lecture_id: int,
    lecture: Lecture = Depends(require_lecture_teacher_access),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Удаление лекции"""
    # Проверка доступа выполнена через зависимость require_lecture_teacher_access
    
    db.delete(lecture)
    db.commit()
    
    return {"message": "Лекция удалена"}


@router.post("/lectures/{lecture_id}/materials")
@limiter.limit("30/minute")
def upload_material(
    request: Request,
    lecture_id: int,
    lecture: Lecture = Depends(require_lecture_teacher_access),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Загрузка материала для лекции"""
    # Проверка доступа выполнена через зависимость require_lecture_teacher_access
    
    # Определяем тип файла
    file_extension = Path(file.filename).suffix.lower()
    file_type_map = {
        '.mp4': 'video', '.avi': 'video', '.mov': 'video', '.mkv': 'video', '.webm': 'video',
        '.pdf': 'pdf',
        '.pptx': 'presentation', '.ppt': 'presentation',
        '.mp3': 'audio', '.wav': 'audio', '.ogg': 'audio', '.m4a': 'audio',
        '.zip': 'scorm', '.scorm': 'scorm'
    }
    file_type = file_type_map.get(file_extension, 'other')
    
    # Определяем максимальный размер файла в зависимости от типа
    MAX_VIDEO_SIZE = 20 * 1024 * 1024  # 20 МБ для видео
    MAX_AUDIO_SIZE = 20 * 1024 * 1024  # 20 МБ для аудио
    MAX_PDF_SIZE = 10 * 1024 * 1024    # 10 МБ для PDF
    MAX_OTHER_SIZE = 50 * 1024 * 1024  # 50 МБ для других типов (SCORM, презентации)
    
    if file_type == 'video':
        max_size = MAX_VIDEO_SIZE
        max_size_mb = 20
    elif file_type == 'audio':
        max_size = MAX_AUDIO_SIZE
        max_size_mb = 20
    elif file_type == 'pdf':
        max_size = MAX_PDF_SIZE
        max_size_mb = 10
    else:
        max_size = MAX_OTHER_SIZE
        max_size_mb = 50
    
    # Проверяем размер файла перед сохранением
    # Читаем файл по частям и проверяем размер
    file_size = 0
    chunk_size = 1024 * 1024  # 1 МБ за раз
    temp_file_path = None
    
    try:
        # Создаем временный файл для проверки размера
        import tempfile
        temp_fd, temp_file_path = tempfile.mkstemp()
        temp_file = os.fdopen(temp_fd, 'wb')
        
        # Сбрасываем позицию файла на начало (на случай если он уже был прочитан)
        file.file.seek(0)
        
        # Читаем файл по частям и проверяем размер
        while True:
            chunk = file.file.read(chunk_size)
            if not chunk:
                break
            
            file_size += len(chunk)
            
            # Проверяем размер на каждой итерации
            if file_size > max_size:
                temp_file.close()
                os.unlink(temp_file_path)
                raise HTTPException(
                    status_code=413,
                    detail=f"Файл слишком большой. Максимальный размер для {file_type}: {max_size_mb} МБ. Размер вашего файла: {file_size / (1024 * 1024):.2f} МБ"
                )
            
            temp_file.write(chunk)
        
        temp_file.close()
        
        # Проверяем финальный размер
        if file_size > max_size:
            os.unlink(temp_file_path)
            raise HTTPException(
                status_code=413,
                detail=f"Файл слишком большой. Максимальный размер для {file_type}: {max_size_mb} МБ. Размер вашего файла: {file_size / (1024 * 1024):.2f} МБ"
            )
        
        # Создаем директорию для лекции
        lecture_dir = UPLOAD_DIR / str(lecture_id)
        lecture_dir.mkdir(parents=True, exist_ok=True)
        
        # Перемещаем временный файл в финальное место
        file_path = lecture_dir / file.filename
        shutil.move(temp_file_path, file_path)
        temp_file_path = None  # Уже перемещен
        
    except HTTPException:
        # Перебрасываем HTTPException
        raise
    except Exception as e:
        # Очищаем временный файл при ошибке
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
            except:
                pass
        logger.error(f"Ошибка при загрузке файла {file.filename}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Ошибка при загрузке файла: {str(e)}")
    
    # Получаем текущий максимальный order_index
    max_order = db.query(LectureMaterial).filter(LectureMaterial.lecture_id == lecture_id).count()
    
    # Сохраняем относительный путь для доступа через статику
    # Формируем путь как uploads/lectures/{lecture_id}/{filename}
    relative_path_str = f"uploads/lectures/{lecture_id}/{file.filename}"
    
    # Создаем запись о материале
    # Используем уже вычисленный file_size вместо повторного чтения размера файла
    material = LectureMaterial(
        lecture_id=lecture_id,
        file_path=relative_path_str,
        file_type=file_type,
        file_name=file.filename,
        file_size=file_size,  # Используем размер, вычисленный при проверке
        order_index=max_order
    )
    
    db.add(material)
    db.commit()
    db.refresh(material)
    
    return {
        "id": material.id,
        "file_path": material.file_path,
        "file_type": material.file_type,
        "file_name": material.file_name,
        "file_size": material.file_size,
        "order_index": material.order_index
    }


@router.delete("/lectures/{lecture_id}/materials/{material_id}")
def delete_material(
    lecture_id: int,
    material_id: int,
    lecture: Lecture = Depends(require_lecture_teacher_access),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Удаление материала лекции"""
    # Проверка доступа выполнена через зависимость require_lecture_teacher_access
    
    material = db.query(LectureMaterial).filter(
        LectureMaterial.id == material_id,
        LectureMaterial.lecture_id == lecture_id
    ).first()
    
    if not material:
        raise HTTPException(status_code=404, detail="Материал не найден")
    
    # Удаляем файл (используем абсолютный путь)
    file_path = Path(material.file_path)
    if not file_path.is_absolute():
        file_path = Path.cwd() / file_path
    if file_path.exists():
        file_path.unlink()
    
    db.delete(material)
    db.commit()
    
    return {"message": "Материал удален"}


@router.put("/lectures/{lecture_id}/materials/reorder")
def reorder_materials(
    lecture_id: int,
    material_ids: List[int],
    lecture: Lecture = Depends(require_lecture_teacher_access),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Изменение порядка материалов"""
    # Проверка доступа выполнена через зависимость require_lecture_teacher_access
    
    # Обновляем порядок
    for order_index, material_id in enumerate(material_ids):
        material = db.query(LectureMaterial).filter(
            LectureMaterial.id == material_id,
            LectureMaterial.lecture_id == lecture_id
        ).first()
        if material:
            material.order_index = order_index
    
    db.commit()
    return {"message": "Порядок материалов обновлен"}


@router.get("/materials/{material_id}/content")
def get_material_content(
    material_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Получение текстового содержимого материала (PDF, Word и т.д.)"""
    # Проверяем доступ через зависимость
    material, lecture, course = require_material_access(
        material_id,
        db,
        current_user,
        require_published=(current_user.role == "student")
    )
    
    # Проверяем, опубликована ли лекция (для админа пропускаем проверку)
    if current_user.role != "admin" and not lecture.published:
        return JSONResponse({
            "content": None,
            "error": "Текст документа доступен только после публикации лекции. Нажмите кнопку 'Выложить' в конструкторе лекции.",
            "file_name": material.file_name,
            "file_type": material.file_type,
            "not_published": True
        })
    
    # Используем ProcessedMaterial (для всех, если лекция опубликована)
    processed_material = db.query(ProcessedMaterial).filter(
        ProcessedMaterial.material_id == material_id
    ).first()
    
    if processed_material and processed_material.processed_text:
        # Возвращаем сохраненный текст из ProcessedMaterial
        return JSONResponse({
            "content": processed_material.processed_text,
            "error": None,
            "file_name": material.file_name,
            "file_type": material.file_type,
            "from_cache": True
        })
    else:
        # Если ProcessedMaterial нет, возвращаем ошибку
        return JSONResponse({
            "content": None,
            "error": "Материал еще не обработан. Обратитесь к преподавателю для публикации лекции.",
            "file_name": material.file_name,
            "file_type": material.file_type
        })


def parse_pdf(file_path: Path) -> str:
    """Парсинг PDF файла и извлечение текста"""
    text_content = []
    
    # Пробуем использовать pdfplumber (более точный)
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(f"--- Страница {i} ---\n\n{page_text}")
            return "\n\n".join(text_content) if text_content else "Текст не найден в PDF"
        except Exception as e:
            print(f"Ошибка pdfplumber: {e}")
    
    # Fallback на PyPDF2
    if PDF_PARSER_AVAILABLE:
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for i, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(f"--- Страница {i} ---\n\n{page_text}")
            return "\n\n".join(text_content) if text_content else "Текст не найден в PDF"
        except Exception as e:
            print(f"Ошибка PyPDF2: {e}")
    
    raise Exception("Библиотеки для парсинга PDF не установлены")


def parse_docx(file_path: Path) -> str:
    """Парсинг Word документа и извлечение текста"""
    if not DOCX_PARSER_AVAILABLE:
        raise Exception("Библиотека python-docx не установлена")
    
    try:
        doc = Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        return "\n\n".join(paragraphs) if paragraphs else "Текст не найден в документе"
    except Exception as e:
        raise Exception(f"Ошибка парсинга Word: {str(e)}")


@router.get("/materials/{material_id}/transcribe")
def transcribe_video(
    material_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Получение транскрипта видео/аудио (из БД для опубликованных лекций)"""
    # Проверяем доступ через зависимость
    material, lecture, course = require_material_access(
        material_id,
        db,
        current_user,
        require_published=(current_user.role == "student")
    )
    
    # Проверяем, что это видео или аудио
    video_exts = ['.mp4', '.avi', '.mov', '.mkv', '.flv', '.wmv', '.webm', '.m4v', '.3gp']
    audio_exts = ['.mp3', '.wav', '.flac', '.aac', '.m4a', '.ogg', '.wma', '.opus']
    file_ext = Path(material.file_name).suffix.lower()
    
    if file_ext not in video_exts and file_ext not in audio_exts:
        raise HTTPException(status_code=400, detail="Транскрибация доступна только для видео и аудио файлов")
    
    # Определяем, является ли пользователь преподавателем
    is_teacher = current_user.role in ["teacher", "admin"]
    
    # Проверяем сохраненный транскрипт из ProcessedMaterial
    processed_material = db.query(ProcessedMaterial).filter(
        ProcessedMaterial.material_id == material_id
    ).first()
    
    # Если транскрипт есть в БД - возвращаем его
    if processed_material and processed_material.processed_text:
        return JSONResponse({
            "text": processed_material.processed_text,
            "language": "ru"
        })
    
    # Если транскрипта нет в БД
    if is_teacher:
        # Для преподавателей: транскрибация доступна только после публикации
        if not lecture.published:
            raise HTTPException(
                status_code=403, 
                detail="Транскрибация доступна только после публикации лекции. Нажмите кнопку 'Выложить' в конструкторе лекции."
            )
        else:
            # Лекция опубликована, но транскрипта нет - значит при публикации произошла ошибка
            raise HTTPException(
                status_code=404, 
                detail="Транскрипт еще не готов. Перепубликуйте лекцию, нажав кнопку 'Выложить' в конструкторе."
            )
    else:
        # Для студентов: транскрипт должен быть в БД после публикации
        raise HTTPException(
            status_code=404, 
            detail="Транскрипт еще не готов. Обратитесь к преподавателю для публикации лекции."
        )




@router.get("/materials/{material_id}/file")
def get_material_file(
    material_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Получение файла материала с проверкой прав доступа"""
    from fastapi.responses import FileResponse
    
    # Проверяем доступ через зависимость
    material, lecture, course = require_material_access(
        material_id,
        db,
        current_user,
        require_published=False  # Файлы доступны даже для неопубликованных лекций (для преподавателей)
    )
    
    # Проверяем существование файла
    file_path = Path(material.file_path)
    if not file_path.is_absolute():
        file_path = Path.cwd() / file_path
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Файл не найден")
    
    # Определяем media type
    media_type_map = {
        'video': 'video/mp4',
        'pdf': 'application/pdf',
        'audio': 'audio/mpeg',
        'presentation': 'application/vnd.openxmlformats-officedocument.presentationml.presentation'
    }
    media_type = media_type_map.get(material.file_type, 'application/octet-stream')
    
    return FileResponse(
        path=str(file_path),
        media_type=media_type,
        filename=material.file_name,
        headers={
            "Content-Disposition": f'inline; filename="{material.file_name}"'
        }
    )


def process_single_material(material: LectureMaterial, lecture_id: int, user_id: int) -> tuple[bool, Optional[str], Optional[list], Optional[str]]:
    """
    Обрабатывает один материал синхронно.
    Возвращает: (success, processed_text, embedding, error_message)
    """
    from app.core.database import SessionLocal
    from app.utils.embeddings import generate_embedding
    
    db = SessionLocal()
    try:
        # Проверяем, не обработан ли уже этот материал
        existing = db.query(ProcessedMaterial).filter(
            ProcessedMaterial.material_id == material.id
        ).first()
        
        if existing:
            # Материал уже обработан
            return (True, existing.processed_text, None, None)
        
        # Получаем путь к файлу
        file_path = Path(material.file_path)
        if not file_path.is_absolute():
            file_path = Path.cwd() / file_path
        
        if not file_path.exists():
            return (False, None, None, f"Файл не найден: {material.file_name}")
        
        # Формируем URL файла
        file_url = f"/api/materials/{material.id}/file"
        
        processed_text = None
        
        # Обработка в зависимости от типа файла
        video_exts = ['.mp4', '.avi', '.mov', '.mkv', '.flv', '.wmv', '.webm', '.m4v', '.3gp']
        audio_exts = ['.mp3', '.wav', '.flac', '.aac', '.m4a', '.ogg', '.wma', '.opus']
        file_ext = Path(material.file_name).suffix.lower()
        
        # Обновляем file_type материала
        if file_ext in video_exts:
            if material.file_type != 'video':
                material.file_type = 'video'
                db.commit()
        elif file_ext in audio_exts:
            if material.file_type != 'audio':
                material.file_type = 'audio'
                db.commit()
        
        if file_ext in video_exts or file_ext in audio_exts:
            # Транскрибация видео или аудио
            logger.info(f"Обработка видео/аудио файла: {material.file_name}")
            if not WHISPER_AVAILABLE:
                return (False, None, None, f"Whisper не установлен, пропущено: {material.file_name}")
            
            try:
                ffmpeg_path = find_ffmpeg()
                if not ffmpeg_path:
                    return (False, None, None, f"FFmpeg не найден, пропущено: {material.file_name}")
                
                ensure_ffmpeg_in_path(ffmpeg_path)
                
                # Транскрибируем синхронно (это уже выполняется в фоновом потоке)
                from app.core.config import WHISPER_MODEL
                processed_text = transcribe_file(file_path, WHISPER_MODEL)
                logger.info(f"Транскрибация завершена: {material.file_name}, длина текста: {len(processed_text) if processed_text else 0} символов")
                
            except Exception as e:
                logger.error(f"Ошибка транскрибации {material.file_name}: {e}", exc_info=True)
                return (False, None, None, f"Ошибка транскрибации {material.file_name}: {str(e)}")
        
        elif material.file_type == 'pdf':
            # Парсинг PDF синхронно
            try:
                import pdfplumber
                text_parts = []
                with pdfplumber.open(str(file_path)) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                processed_text = "\n\n".join(text_parts)
                logger.info(f"Распарсен PDF: {material.file_name}")
                
            except Exception as e:
                logger.error(f"Ошибка парсинга PDF {material.file_name}: {e}")
                return (False, None, None, f"Ошибка парсинга PDF {material.file_name}: {str(e)}")
        
        elif material.file_name.endswith('.docx') or material.file_name.endswith('.doc'):
            # Парсинг DOCX синхронно
            try:
                from docx import Document
                doc = Document(str(file_path))
                text_parts = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                processed_text = "\n\n".join(text_parts)
                logger.info(f"Распарсен DOCX: {material.file_name}")
                
            except Exception as e:
                logger.error(f"Ошибка парсинга DOCX {material.file_name}: {e}")
                return (False, None, None, f"Ошибка парсинга DOCX {material.file_name}: {str(e)}")
        
        # Генерируем эмбеддинг для текста, если он есть
        embedding = None
        if processed_text and processed_text.strip():
            try:
                text_for_embedding = processed_text.strip()
                
                # Генерируем эмбеддинг синхронно
                embedding = generate_embedding(
                    text_for_embedding,
                    use_chunks=True,
                    chunk_size=2000,
                    overlap=200
                )
                
                if embedding:
                    logger.info(f"Сгенерирован эмбеддинг для материала: {material.file_name}")
                else:
                    logger.warning(f"Эмбеддинг не был сгенерирован для {material.file_name}")
            except Exception as e:
                logger.error(f"Ошибка генерации эмбеддинга для {material.file_name}: {e}", exc_info=True)
        
        # Сохраняем обработанный материал
        processed_material = ProcessedMaterial(
            lecture_id=lecture_id,
            material_id=material.id,
            user_id=user_id,
            file_url=file_url,
            file_type=material.file_type,
            processed_text=processed_text,
            embedding=embedding,
            processed_at=datetime.now().isoformat()
        )
        db.add(processed_material)
        db.commit()
        logger.info(f"Обработан материал: {material.file_name} (тип: {material.file_type})")
        
        return (True, processed_text, embedding, None)
        
    except Exception as e:
        logger.error(f"Ошибка обработки материала {material.file_name}: {e}", exc_info=True)
        db.rollback()
        return (False, None, None, f"Ошибка обработки {material.file_name}: {str(e)}")
    finally:
        db.close()


def process_lecture_materials_background(lecture_id: int, user_id: int):
    """
    Фоновая обработка материалов лекции.
    Выполняется в отдельном потоке, чтобы не блокировать основной event loop.
    """
    from app.core.database import SessionLocal
    from concurrent.futures import ThreadPoolExecutor, as_completed
    
    db = SessionLocal()
    try:
        lecture = db.query(Lecture).filter(Lecture.id == lecture_id).first()
        if not lecture:
            logger.error(f"Лекция {lecture_id} не найдена для фоновой обработки")
            return
        
        # Получаем все материалы лекции
        materials = db.query(LectureMaterial).filter(LectureMaterial.lecture_id == lecture_id).all()
        
        if not materials:
            logger.warning(f"Лекция {lecture_id} не содержит материалов")
            return
        
        logger.info(f"Начало фоновой обработки {len(materials)} материалов для лекции {lecture_id}")
        
        # Обрабатываем материалы параллельно используя ThreadPoolExecutor
        processed_count = 0
        errors = []
        
        # Используем ThreadPoolExecutor для параллельной обработки
        with ThreadPoolExecutor(max_workers=min(len(materials), 4)) as executor:
            # Запускаем обработку всех материалов
            future_to_material = {
                executor.submit(process_single_material, material, lecture_id, user_id): material
                for material in materials
            }
            
            # Собираем результаты
            for future in as_completed(future_to_material):
                material = future_to_material[future]
                try:
                    result = future.result()
                    if isinstance(result, tuple):
                        success, _, _, error = result
                        if success:
                            processed_count += 1
                        elif error:
                            errors.append(error)
                except Exception as e:
                    errors.append(f"Ошибка обработки {material.file_name}: {str(e)}")
                    logger.error(f"Исключение при обработке материала {material.file_name}: {e}", exc_info=True)
        
        # Публикуем лекцию ТОЛЬКО если ВСЕ материалы успешно обработаны
        db_refresh = SessionLocal()
        try:
            lecture_refresh = db_refresh.query(Lecture).filter(Lecture.id == lecture_id).first()
            if not lecture_refresh:
                logger.error(f"Лекция {lecture_id} не найдена при публикации")
                return
            
            if processed_count == len(materials) and len(errors) == 0:
                # Используем транзакцию для атомарности: публикация лекции + генерация теста
                try:
                    # Начинаем транзакцию
                    lecture_refresh.published = True
                    logger.info(f"Публикация лекции {lecture_id}. Обработано: {processed_count}/{len(materials)}")
                    
                    # Генерируем тест, если нужно (в той же транзакции)
                    if lecture_refresh.generate_test and lecture_refresh.test_generation_mode == "once":
                        from app.utils.rag import generate_questions_from_text
                        
                        logger.info(f"Начинаем генерацию теста для лекции {lecture_id}")
                        
                        processed_materials = db_refresh.query(ProcessedMaterial).filter(
                            ProcessedMaterial.lecture_id == lecture_id,
                            ProcessedMaterial.processed_text.isnot(None)
                        ).all()
                        
                        all_questions = []
                        order_index = 0
                        
                        for pm in processed_materials:
                            if not pm.processed_text or not pm.processed_text.strip():
                                continue
                            
                            text = pm.processed_text.strip()
                            text_length = len(text)
                            
                            if text_length < 500:
                                num_questions = 2
                            elif text_length < 1500:
                                num_questions = 2 if text_length < 1000 else 3
                            else:
                                num_questions = 3
                            
                            questions_data = generate_questions_from_text(text, num_questions=num_questions)
                            
                            if questions_data and len(questions_data) > 0:
                                for q_data in questions_data:
                                    q_data["order_index"] = order_index
                                    order_index += 1
                                all_questions.extend(questions_data)
                        
                        if all_questions and len(all_questions) > 0:
                            test = Test(
                                lecture_id=lecture_id,
                                created_at=datetime.now().isoformat()
                            )
                            db_refresh.add(test)
                            db_refresh.flush()
                            
                            for q_data in all_questions:
                                question = Question(
                                    test_id=test.id,
                                    question_text=q_data["question_text"],
                                    correct_answer=q_data["correct_answer"],
                                    options=q_data.get("options"),
                                    question_type=q_data["question_type"],
                                    order_index=q_data["order_index"]
                                )
                                db_refresh.add(question)
                        
                        logger.info(f"✅ Подготовлен тест из {len(all_questions)} вопросов для лекции {lecture_id}")
                    
                    # Коммитим все изменения атомарно
                    db_refresh.commit()
                    logger.info(f"Лекция {lecture_id} успешно опубликована. Обработано: {processed_count}/{len(materials)}")
                except Exception as e:
                    # Откатываем транзакцию при любой ошибке
                    db_refresh.rollback()
                    logger.error(f"❌ Ошибка при публикации лекции {lecture_id}: {e}. Транзакция откачена.", exc_info=True)
                    raise
            else:
                logger.warning(f"Не удалось обработать все материалы для лекции {lecture_id}. Обработано: {processed_count}/{len(materials)}, ошибки: {errors}")
        finally:
            db_refresh.close()
        
    except Exception as e:
        logger.error(f"Критическая ошибка при фоновой обработке лекции {lecture_id}: {e}", exc_info=True)
    finally:
        db.close()


@router.post("/lectures/{lecture_id}/publish")
@limiter.limit("10/hour")
def publish_lecture(
    request: Request,
    lecture_id: int,
    background_tasks: BackgroundTasks,
    lecture: Lecture = Depends(require_lecture_teacher_access),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Публикация лекции: запускает фоновую обработку материалов (транскрибация, парсинг, эмбеддинги)"""
    logger.info(f"Запрос на публикацию лекции {lecture_id} от пользователя {current_user.id}")
    
    # Проверка доступа выполнена через зависимость require_lecture_teacher_access
    
    # Получаем все материалы лекции
    materials = db.query(LectureMaterial).filter(LectureMaterial.lecture_id == lecture_id).all()
    
    if not materials:
        raise HTTPException(status_code=400, detail="Лекция не содержит материалов")
    
    # Проверяем, не запущена ли уже обработка
    if lecture.published:
        return JSONResponse({
            "message": "Лекция уже опубликована",
            "published": True,
            "processing": False
        })
    
    # Запускаем фоновую обработку материалов
    background_tasks.add_task(process_lecture_materials_background, lecture_id, current_user.id)
    
    logger.info(f"Запущена фоновая обработка материалов для лекции {lecture_id}")
    
    return JSONResponse({
        "message": "Обработка материалов начата. Лекция будет опубликована после завершения обработки всех материалов.",
        "lecture_id": lecture_id,
        "materials_count": len(materials),
        "processing": True,
        "published": False
    })
